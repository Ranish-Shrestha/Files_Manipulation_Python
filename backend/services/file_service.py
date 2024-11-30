import os
import docx
from pypdf import PdfReader
from bson import ObjectId
from datetime import datetime
from pymongo import MongoClient
from flask import jsonify, send_file
from models.file_model import file_collection, log_collection, file_schema, log_schema

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = "".join(page.extract_text() for page in reader.pages)
    return text

def read_docx(file_path):
    doc = docx.Document(file_path)
    text = "\n".join(para.text for para in doc.paragraphs)
    return text

def log_activity(action, details):
    log_collection.insert_one(log_schema(action, details))

def upload_file_service(request):
    if 'files' not in request.files:
        return jsonify({"error": "No files part"}), 400

    files = request.files.getlist('files')
    if not files:
        return jsonify({"error": "No selected files"}), 400

    responses = []

    for file in files:
        if file.filename == '':
            continue

        file_path = os.path.join("uploads", file.filename)
        file.save(file_path)

        if file.filename.endswith('.pdf'):
            text = read_pdf(file_path)
        elif file.filename.endswith('.docx'):
            text = read_docx(file_path)
        else:
            responses.append({"filename": file.filename, "status": "Unsupported file type"})
            continue

        file_record = file_schema(file.filename, os.path.splitext(file.filename)[1], os.path.getsize(file_path), text)
        file_collection.insert_one(file_record)
        log_activity("upload", {"filename": file.filename})
        responses.append({"filename": file.filename, "status": "File uploaded and content stored"})

    return jsonify(responses), 200

def get_files_service():
    files = list(file_collection.find({}, {'content': 0}))
    return jsonify(files), 200

def generate_pdf(content, filename):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in content.split("\n"):
        pdf.cell(200, 10, txt=line, ln=True)
    output_path = os.path.join("downloads", filename + ".pdf")
    pdf.output(output_path)
    return output_path

def generate_docx(content, filename):
    doc = docx.Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    output_path = os.path.join("downloads", filename + ".docx")
    doc.save(output_path)
    return output_path

def download_file_service(file_id, file_type):
    file_record = file_collection.find_one({"_id": ObjectId(file_id)})
    if not file_record:
        return jsonify({"error": "File not found"}), 404

    filename = file_record["filename"]
    content = file_record["content"]

    if file_type == "pdf":
        file_path = generate_pdf(content, filename)
    elif file_type == "docx":
        file_path = generate_docx(content, filename)
    else:
        return jsonify({"error": "Unsupported file type"}), 400

    file_collection.update_one({"_id": ObjectId(file_id)}, {"$inc": {"download_count": 1}})
    log_activity("download", {"filename": file_record["filename"], "file_type": file_type})
    return send_file(file_path, as_attachment=True)

def get_activity_log_service():
    logs = list(log_collection.find())
    return jsonify(logs), 200
