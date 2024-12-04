import io
import docx
from pypdf import PdfReader
from fpdf import FPDF
from fastapi import HTTPException, UploadFile


async def read_pdf(file: UploadFile):
    try:
        file.file.seek(0)
        buffer = io.BytesIO(file.file.read())
        pdf_reader = PdfReader(buffer)
        if pdf_reader.is_encrypted:
            raise HTTPException(status_code=400, detail="PDF file is encrypted and cannot be read")
        text = "".join(page.extract_text() for page in pdf_reader.pages if page.extract_text())
        if not text.strip():
            raise HTTPException(status_code=400, detail="PDF file is empty")
        return text
    except Exception:
        raise HTTPException(status_code=400, detail="Unable to read PDF file")

async def read_docx(file: UploadFile):
    try:
        file.file.seek(0)
        buffer = io.BytesIO(file.file.read())
        doc_reader = docx.Document(buffer)
        text = "\n".join(para.text for para in doc_reader.paragraphs)
        if not text.strip():
            raise HTTPException(status_code=400, detail="DOCX file is empty")
        return text
    except docx.opc.exceptions.PackageNotFoundError:
        raise HTTPException(status_code=400, detail="DOCX file is corrupted or encrypted and cannot be read")
    except Exception:
        raise HTTPException(status_code=400, detail="Unable to read DOCX file")

async def generate_pdf(content):
    pdf_writer = FPDF()
    pdf_writer.add_page()

    # Set font
    pdf_writer.add_font('DejaVu', '', 'fonts/DejaVuSans.ttf', uni=True)
    pdf_writer.set_font('DejaVu', '', 10)

    # Set margins: left, top, right
    pdf_writer.set_left_margin(10)
    pdf_writer.set_top_margin(10)
    pdf_writer.set_right_margin(10)
    pdf_writer.set_auto_page_break(auto=True, margin=15)

    # Apply the margins
    pdf_writer.set_xy(pdf_writer.l_margin, pdf_writer.t_margin)

    for line in content.split("\n"):
        pdf_writer.multi_cell(0, 10, txt=line, ln=True)

    buffer = io.BytesIO()
    pdf_writer.output(buffer)
    buffer.seek(0)
    return buffer

async def generate_docx(content):
    doc_writer = docx.Document()
    
    for line in content.split("\n"):
        doc_writer.add_paragraph(line)

    buffer = io.BytesIO()
    doc_writer.save(buffer)
    buffer.seek(0)
    return buffer
